import io
import os
import re
from datetime import datetime
from html import escape
from typing import Any, List, Optional


def _format_date(value: Any) -> str:
    if value is None:
        return "Not recorded"
    if isinstance(value, str):
        try:
            value = datetime.fromisoformat(value)
        except ValueError:
            return value
    return value.strftime("%d %B %Y")


def _org(organization: Optional[dict]) -> dict:
    data = organization or {}
    return {
        "name": data.get("name") or "Starlight Secondary School",
        "address": data.get("address") or "",
        "phone": data.get("phone") or "",
        "email": data.get("email") or "",
        "website": data.get("website") or "",
        "motto": data.get("motto") or "",
        "registration": data.get("registration") or "",
        "logo_path": data.get("logo_path") or "",
    }


def _usable_logo(path: str) -> bool:
    return bool(path and os.path.isfile(path) and not path.lower().endswith(".svg"))


def _participant_rows(participants: List[Any]) -> List[dict]:
    rows = []
    for participant in participants:
        if isinstance(participant, dict):
            rows.append({
                "name": participant.get("name", ""),
                "role": participant.get("role", ""),
                "email": participant.get("email", ""),
                "status": participant.get("status", "Present"),
            })
        else:
            rows.append({"name": str(participant), "role": "", "email": "", "status": "Present"})
    return rows


def _agenda_items(description: str) -> List[str]:
    text = (description or "").strip()
    if not text:
        return ["Agenda not recorded"]
    items = [
        re.sub(r"^\s*(?:\d+[\.\)]|[-•])\s*", "", item).strip()
        for item in re.split(r"(?:\r?\n|;|,\s*)+", text)
        if item.strip()
    ]
    if len(items) == 1:
        detected = [
            item.strip()
            for item in re.split(
                r"(?i)(?=\b(?:opening prayer|prayer|introduction(?:s)?|"
                r"communication from (?:the )?chair(?:person)?|reactions?|"
                r"way\s*forward|matters arising|any other business|closing)\b)",
                text,
            )
            if item.strip()
        ]
        if len(detected) > 1:
            items = detected
    return items or [text]


def export_txt(
    meeting_title: str, meeting_date: Any, participants: List[str],
    summary_text: str, decisions: List[str], action_items: List[dict],
    transcript_text: str, organization: Optional[dict] = None,
    meeting_description: str = "",
) -> bytes:
    org = _org(organization)
    contacts = " | ".join(filter(None, [org["address"], org["phone"], org["email"], org["website"]]))
    attendee_rows = _participant_rows(participants)
    agenda = _agenda_items(meeting_description)
    lines = [
        "=" * 78, org["name"].upper(), contacts, org["motto"],
        "OFFICIAL MEETING MINUTES", "=" * 78,
        f"Meeting          : {meeting_title}",
        f"Date             : {_format_date(meeting_date)}",
        "AGENDA",
        "", "MINUTES / EXECUTIVE RECORD", "-" * 45,
        summary_text or "No minutes have been generated.", "", "KEY DECISIONS", "-" * 45,
    ]
    lines[lines.index("AGENDA") + 1:lines.index("AGENDA") + 1] = [
        f"{index}. {item}" for index, item in enumerate(agenda, 1)
    ]
    attendance_index = lines.index("MINUTES / EXECUTIVE RECORD")
    lines[attendance_index:attendance_index] = [
        "", "ATTENDANCE REGISTER", "-" * 45,
        *[
            f"{index}. {row['name']} | {row['role'] or 'Stakeholder'} | "
            f"{row['email'] or 'No email'} | {row['status']}"
            for index, row in enumerate(attendee_rows, 1)
        ],
        "",
    ]
    lines.extend(f"{index}. {text}" for index, text in enumerate(decisions, 1))
    if not decisions:
        lines.append("No decisions recorded.")
    lines.extend(["", "ACTION ITEMS", "-" * 45])
    for index, item in enumerate(action_items, 1):
        lines.append(
            f"{index}. {item.get('text', '')} | Responsible: "
            f"{item.get('assignee_name') or 'Unassigned'} | Due: "
            f"{item.get('deadline') or 'Not specified'} | "
            f"Status: {item.get('status', 'pending').replace('_', ' ').title()}"
        )
    if not action_items:
        lines.append("No action items recorded.")
    lines.extend([
        "", "APPENDIX: FULL TRANSCRIPT", "-" * 45,
        transcript_text or "No transcript available.", "",
        f"Prepared on {datetime.now().strftime('%d %B %Y at %H:%M')}",
    ])
    return "\n".join(filter(lambda line: line is not None, lines)).encode("utf-8")


def export_pdf(
    meeting_title: str, meeting_date: Any, participants: List[str],
    summary_text: str, decisions: List[str], action_items: List[dict],
    transcript_text: str, organization: Optional[dict] = None,
    meeting_description: str = "",
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    org = _org(organization)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, leftMargin=1.7 * cm, rightMargin=1.7 * cm,
        topMargin=1.5 * cm, bottomMargin=1.6 * cm,
        title=f"Minutes - {meeting_title}", author=org["name"],
    )
    styles = getSampleStyleSheet()
    blue = colors.HexColor("#1746A2")
    dark = colors.HexColor("#111827")
    body = ParagraphStyle("MinutesBody", parent=styles["BodyText"], fontSize=10, leading=15, textColor=dark, spaceAfter=6)
    small = ParagraphStyle("MinutesSmall", parent=body, fontSize=8, leading=11, textColor=colors.HexColor("#4b5563"))
    heading = ParagraphStyle("MinutesHeading", parent=styles["Heading2"], fontSize=12, leading=15, textColor=blue, spaceBefore=12, spaceAfter=6)
    org_title = ParagraphStyle("OrgTitle", parent=styles["Title"], fontSize=17, leading=20, textColor=blue, alignment=1, spaceAfter=2)
    centered = ParagraphStyle("Centered", parent=small, alignment=1)
    story = []

    logo_path = org["logo_path"]
    org_text = [
        Paragraph(f"<b>{escape(org['name'].upper())}</b>", org_title),
        Paragraph(escape(" | ".join(filter(None, [org["address"], org["phone"], org["email"], org["website"]]))), centered),
    ]
    if org["motto"]:
        org_text.append(Paragraph(f"<i>{escape(org['motto'])}</i>", centered))
    if _usable_logo(logo_path):
        logo = Image(logo_path, width=2.1 * cm, height=2.1 * cm, kind="proportional")
        header = Table([[logo, org_text]], colWidths=[2.5 * cm, 14.5 * cm])
        header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
        story.append(header)
    else:
        story.extend(org_text)
    if org["registration"]:
        story.append(Paragraph(f"Registration/Centre No: {escape(org['registration'])}", centered))
    story.append(HRFlowable(width="100%", thickness=1.2, color=blue, spaceBefore=5, spaceAfter=9))
    story.append(Paragraph("OFFICIAL MEETING MINUTES", org_title))

    details = [
        ["Meeting", Paragraph(escape(meeting_title), body)],
        ["Date", _format_date(meeting_date)],
    ]
    detail_table = Table(details, colWidths=[3.5 * cm, 13.5 * cm])
    detail_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eaf1ff")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(detail_table)

    story.append(Paragraph("1. Agenda", heading))
    for index, item in enumerate(_agenda_items(meeting_description), 1):
        story.append(Paragraph(f"<b>{index}.</b> {escape(item)}", body))

    story.append(Paragraph("2. Attendance Register", heading))
    attendee_rows = _participant_rows(participants)
    if attendee_rows:
        attendance_table_data = [["No.", "Name", "Role / Organization", "Email", "Status"]]
        for index, row in enumerate(attendee_rows, 1):
            attendance_table_data.append([
                str(index),
                Paragraph(escape(row["name"]), small),
                Paragraph(escape(row["role"] or "Stakeholder"), small),
                Paragraph(escape(row["email"] or "—"), small),
                row["status"],
            ])
        attendance_table = Table(
            attendance_table_data,
            colWidths=[0.8 * cm, 4.2 * cm, 4.3 * cm, 5.2 * cm, 2.5 * cm],
            repeatRows=1,
        )
        attendance_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), blue),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(attendance_table)
    else:
        story.append(Paragraph("No attendance recorded.", body))

    story.append(Paragraph("3. Minutes / Proceedings", heading))
    for paragraph in (summary_text or "No minutes have been generated.").splitlines():
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph.strip()), body))

    story.append(Paragraph("4. Resolutions / Key Decisions", heading))
    for index, decision in enumerate(decisions, 1):
        story.append(Paragraph(f"<b>{index}.</b> {escape(decision)}", body))
    if not decisions:
        story.append(Paragraph("No decisions recorded.", body))

    story.append(Paragraph("5. Action Items", heading))
    if action_items:
        rows = [["#", "Action", "Responsible", "Due date", "Status"]]
        for index, item in enumerate(action_items, 1):
            rows.append([
                str(index), Paragraph(escape(item.get("text", "")), small),
                Paragraph(escape(item.get("assignee_name") or "Unassigned"), small),
                str(item.get("deadline") or "Not specified"),
                str(item.get("status", "pending")).replace("_", " ").title(),
            ])
        table = Table(rows, colWidths=[0.6 * cm, 7.4 * cm, 3.2 * cm, 3.2 * cm, 2.6 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), blue), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"), ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"), ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(table)
    else:
        story.append(Paragraph("No action items recorded.", body))

    story.extend([Spacer(1, 0.3 * cm), Paragraph("Appendix: Full Transcript", heading)])
    for paragraph in (transcript_text or "No transcript available.").splitlines():
        if paragraph.strip():
            story.append(Paragraph(escape(paragraph.strip()), small))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(
        f"Prepared electronically on {datetime.now().strftime('%d %B %Y at %H:%M')}.",
        centered,
    ))

    def footer(canvas, document):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#6b7280"))
        canvas.drawCentredString(A4[0] / 2, 0.8 * cm, f"{org['name']} · Official Minutes · Page {document.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return buffer.getvalue()


def export_docx(
    meeting_title: str, meeting_date: Any, participants: List[str],
    summary_text: str, decisions: List[str], action_items: List[dict],
    transcript_text: str, organization: Optional[dict] = None,
    meeting_description: str = "",
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    org = _org(organization)
    doc = Document()
    blue = RGBColor(0x17, 0x46, 0xA2)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"{org['name']} · Official Meeting Minutes").font.size = Pt(8)

    header = doc.sections[0].header
    table = header.add_table(rows=1, cols=2, width=Inches(6.7))
    logo_path = org["logo_path"]
    if _usable_logo(logo_path):
        table.cell(0, 0).paragraphs[0].add_run().add_picture(logo_path, width=Inches(0.75))
    center = table.cell(0, 1)
    name_p = center.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(org["name"].upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = blue
    contacts = " | ".join(filter(None, [org["address"], org["phone"], org["email"], org["website"]]))
    if contacts:
        p = center.add_paragraph(contacts)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
    if org["motto"]:
        p = center.add_paragraph(org["motto"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    title = doc.add_heading("OFFICIAL MEETING MINUTES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = blue
        run.font.size = Pt(18)

    details = doc.add_table(rows=0, cols=2)
    details.style = "Table Grid"
    for label, value in [
        ("Meeting", meeting_title), ("Date", _format_date(meeting_date)),
    ]:
        cells = details.add_row().cells
        cells[0].text, cells[1].text = label, value
        cells[0].paragraphs[0].runs[0].bold = True

    def add_heading(text: str):
        heading = doc.add_heading(text, level=1)
        for heading_run in heading.runs:
            heading_run.font.color.rgb = blue

    add_heading("1. Agenda")
    for item in _agenda_items(meeting_description):
        doc.add_paragraph(item, style="List Number")

    add_heading("2. Attendance Register")
    attendee_rows = _participant_rows(participants)
    if attendee_rows:
        attendance_table = doc.add_table(rows=1, cols=5)
        attendance_table.style = "Table Grid"
        for cell, label in zip(attendance_table.rows[0].cells, ["No.", "Name", "Role / Organization", "Email", "Status"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for index, row in enumerate(attendee_rows, 1):
            cells = attendance_table.add_row().cells
            for cell, value in zip(cells, [
                str(index), row["name"], row["role"] or "Stakeholder",
                row["email"] or "—", row["status"],
            ]):
                cell.text = value
    else:
        doc.add_paragraph("No attendance recorded.")

    add_heading("3. Minutes / Proceedings")
    for paragraph in (summary_text or "No minutes have been generated.").splitlines():
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    add_heading("4. Resolutions / Key Decisions")
    if decisions:
        for decision in decisions:
            doc.add_paragraph(decision, style="List Number")
    else:
        doc.add_paragraph("No decisions recorded.")
    add_heading("5. Action Items")
    if action_items:
        action_table = doc.add_table(rows=1, cols=5)
        action_table.style = "Table Grid"
        for cell, label in zip(action_table.rows[0].cells, ["#", "Action", "Responsible", "Due date", "Status"]):
            cell.text = label
            cell.paragraphs[0].runs[0].bold = True
        for index, item in enumerate(action_items, 1):
            cells = action_table.add_row().cells
            values = [
                str(index), item.get("text", ""), item.get("assignee_name") or "Unassigned",
                str(item.get("deadline") or "Not specified"),
                str(item.get("status", "pending")).replace("_", " ").title(),
            ]
            for cell, value in zip(cells, values):
                cell.text = value
    else:
        doc.add_paragraph("No action items recorded.")
    add_heading("Appendix: Full Transcript")
    doc.add_paragraph(transcript_text or "No transcript available.")
    generated = doc.add_paragraph(f"Prepared electronically on {datetime.now().strftime('%d %B %Y at %H:%M')}.")
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.runs[0].font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
